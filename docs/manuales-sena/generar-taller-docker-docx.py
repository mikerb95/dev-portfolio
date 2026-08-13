#!/usr/bin/env python3
# Genera "Taller-Docker-Linux-CodeByMike.docx" en formato APA 7 (documento de
# estudiante). El .docx no se edita a mano: se corrige aqui y se regenera, para
# que el documento entregado y su fuente no se separen.
#
#   python3 docs/manuales-sena/generar-taller-docker-docx.py
#
# Requiere python-docx. Los recuadros de figura son huecos por llenar: cada uno
# indica donde tomar la captura, que comandos ejecutar y que debe verse.
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/mike/dev/work/github.com/portfolio/docs/manuales-sena/Taller-Docker-Linux-CodeByMike.docx"

doc = Document()

# ---------------------------------------------------------------- estilos base
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = st.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.first_line_indent = Cm(1.27)

for s in doc.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.54)

# Encabezados APA en Times New Roman negro (Word los trae en azul Calibri)
for name, size, bold, italic, align in [
    ("Heading 1", 12, True, False, WD_ALIGN_PARAGRAPH.CENTER),
    ("Heading 2", 12, True, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("Heading 3", 12, True, True, WD_ALIGN_PARAGRAPH.LEFT),
]:
    h = doc.styles[name]
    h.font.name = "Times New Roman"
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.italic = italic
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.alignment = align
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(0)
    h.paragraph_format.first_line_indent = Cm(0)
    h.paragraph_format.keep_with_next = True
    # Word hereda la fuente del tema (majorHAnsi, sans serif) y esa referencia
    # gana sobre w:ascii. Hay que retirarla para que el titulo salga en Times.
    rf = h.element.rPr.rFonts
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        if rf.get(qn("w:" + attr)) is not None:
            del rf.attrib[qn("w:" + attr)]
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:cs"), "Times New Roman")

lb = doc.styles["List Bullet"]
lb.font.name = "Times New Roman"
lb.font.size = Pt(12)


def shade(cell_or_par, hexcolor):
    el = cell_or_par._tc if hasattr(cell_or_par, "_tc") else cell_or_par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    el.append(sh)


def borders(tbl, top=True, bottom=True, inside=False, color="000000", sz=6):
    """Bordes APA: solo horizontales, nunca verticales."""
    tblPr = tbl._tbl.tblPr
    el = OxmlElement("w:tblBorders")
    conf = {
        "top": top, "bottom": bottom, "left": False, "right": False,
        "insideH": inside, "insideV": False,
    }
    for edge, on in conf.items():
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single" if on else "none")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        el.append(e)
    tblPr.append(el)


def p(text="", *, align=None, indent=True, space_after=0, bold=False,
      italic=False, size=12, spacing=WD_LINE_SPACING.DOUBLE, font="Times New Roman"):
    par = doc.add_paragraph()
    par.paragraph_format.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.first_line_indent = Cm(1.27) if indent else Cm(0)
    par.paragraph_format.line_spacing_rule = spacing
    par.paragraph_format.space_after = Pt(space_after)
    if text:
        r = par.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = font
    return par


def rich(par, parts):
    """parts: lista de (texto, {'bold':..,'italic':..,'mono':..})"""
    for txt, fmt in parts:
        r = par.add_run(txt)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if fmt.get("mono"):
            r.font.name = "Courier New"
            r.font.size = Pt(10.5)
    return par


def code(lines, comment=None):
    """Bloque de comandos: Courier New, interlineado sencillo, fondo gris."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, "F2F2F2")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for i, ln in enumerate(lines):
        cp = cell.add_paragraph()
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cp.paragraph_format.first_line_indent = Cm(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.space_before = Pt(4 if i == 0 else 0)
        r = cp.add_run(ln)
        r.font.name = "Courier New"
        r.font.size = Pt(10)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(4)
    borders(tbl, top=False, bottom=False)
    p("", space_after=0)
    return tbl


def bullets(items):
    for it in items:
        par = doc.add_paragraph(style="List Bullet")
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.left_indent = Cm(1.27)
        if isinstance(it, str):
            r = par.add_run(it)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        else:
            rich(par, it)
            for r in par.runs:
                if r.font.name != "Courier New":
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)


TABLE_N = [0]
FIG_N = [0]


def table(title, headers, rows, note=None, widths=None, fontsize=10):
    """Tabla en formato APA 7."""
    TABLE_N[0] += 1
    n = TABLE_N[0]
    t1 = p(f"Tabla {n}", indent=False, bold=True)
    t1.paragraph_format.keep_with_next = True
    t2 = p(title, indent=False, italic=True, space_after=4)
    t2.paragraph_format.keep_with_next = True

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(fontsize)
        r.font.name = "Times New Roman"
        hdr[i].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        hdr[i].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(2)
        hdr[i].paragraphs[0].paragraph_format.space_before = Pt(2)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            par = cells[i].paragraphs[0]
            par.text = ""
            par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            par.paragraph_format.first_line_indent = Cm(0)
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.space_before = Pt(2)
            if isinstance(val, str):
                r = par.add_run(val)
                r.font.size = Pt(fontsize)
                r.font.name = "Times New Roman"
            else:
                for txt, fmt in val:
                    r = par.add_run(txt)
                    r.bold = fmt.get("bold", False)
                    r.italic = fmt.get("italic", False)
                    r.font.size = Pt(fontsize - 0.5 if fmt.get("mono") else fontsize)
                    r.font.name = "Courier New" if fmt.get("mono") else "Times New Roman"
    borders(tbl, top=True, bottom=True, inside=False)
    # línea bajo el encabezado
    hp = tbl.rows[0].cells[0]._tc.get_or_add_tcPr()
    for c in tbl.rows[0].cells:
        tcPr = c._tc.get_or_add_tcPr()
        b = OxmlElement("w:tcBorders")
        bo = OxmlElement("w:bottom")
        bo.set(qn("w:val"), "single")
        bo.set(qn("w:sz"), "6")
        bo.set(qn("w:color"), "000000")
        b.append(bo)
        tcPr.append(b)
    if widths:
        for r_ in tbl.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    np = p("", indent=False, space_after=0)
    if note:
        np.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        np.paragraph_format.space_before = Pt(4)
        rich(np, [("Nota. ", {"italic": True})] + note)
        for r in np.runs:
            if r.font.name != "Courier New":
                r.font.name = "Times New Roman"
            r.font.size = Pt(r.font.size or Pt(10)) if r.font.size else Pt(10)
        for r in np.runs:
            if r.font.name == "Times New Roman":
                r.font.size = Pt(10)
    p("", space_after=0)
    return tbl


def figura(titulo, donde, comandos, que_debe_verse, altura_cm=7.0):
    """Placeholder de captura en formato de figura APA 7."""
    FIG_N[0] += 1
    n = FIG_N[0]
    f1 = p(f"Figura {n}", indent=False, bold=True)
    f1.paragraph_format.keep_with_next = True
    f2 = p(titulo, indent=False, italic=True, space_after=4)
    f2.paragraph_format.keep_with_next = True

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, "FAFAFA")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)

    head = cell.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    head.paragraph_format.first_line_indent = Cm(0)
    head.paragraph_format.space_before = Pt(8)
    head.paragraph_format.space_after = Pt(6)
    r = head.add_run(f"[ AQUÍ VA LA CAPTURA {n} ]")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

    lbl = cell.add_paragraph()
    lbl.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lbl.paragraph_format.first_line_indent = Cm(0)
    lbl.paragraph_format.space_after = Pt(2)
    rr = lbl.add_run("Donde: ")
    rr.bold = True
    rr.font.size = Pt(9.5)
    rr.font.name = "Times New Roman"
    rr = lbl.add_run(donde)
    rr.font.size = Pt(9.5)
    rr.font.name = "Times New Roman"

    lbl2 = cell.add_paragraph()
    lbl2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lbl2.paragraph_format.first_line_indent = Cm(0)
    lbl2.paragraph_format.space_after = Pt(2)
    rr = lbl2.add_run("Comandos a ejecutar:")
    rr.bold = True
    rr.font.size = Pt(9.5)
    rr.font.name = "Times New Roman"

    for c in comandos:
        cp = cell.add_paragraph()
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cp.paragraph_format.first_line_indent = Cm(0)
        cp.paragraph_format.left_indent = Cm(0.6)
        cp.paragraph_format.space_after = Pt(0)
        rr = cp.add_run(c)
        rr.font.name = "Courier New"
        rr.font.size = Pt(9)

    lbl3 = cell.add_paragraph()
    lbl3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lbl3.paragraph_format.first_line_indent = Cm(0)
    lbl3.paragraph_format.space_before = Pt(4)
    lbl3.paragraph_format.space_after = Pt(10)
    rr = lbl3.add_run("Debe verse: ")
    rr.bold = True
    rr.font.size = Pt(9.5)
    rr.font.name = "Times New Roman"
    rr = lbl3.add_run(que_debe_verse)
    rr.font.size = Pt(9.5)
    rr.font.name = "Times New Roman"

    # El recuadro no debe partirse entre dos paginas: la evidencia se pega
    # dentro de el y una figura cortada arruina la maquetacion.
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)
    for cp in cell.paragraphs[:-1]:
        cp.paragraph_format.keep_with_next = True

    # borde punteado para que se note que es un hueco por llenar
    tblPr = tbl._tbl.tblPr
    el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "dashed")
        e.set(qn("w:sz"), "8")
        e.set(qn("w:color"), "999999")
        el.append(e)
    tblPr.append(el)

    np = p("", indent=False, space_after=0)
    np.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    np.paragraph_format.space_before = Pt(4)
    rich(np, [("Nota. ", {"italic": True}),
              ("Captura tomada en Ubuntu 24.04.4 LTS con Docker Engine 29.6.2. "
               "Elaboración propia.", {})])
    for r in np.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
    p("", space_after=0)


def toc():
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(0)
    run = par.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Actualice la tabla de contenido en Word: clic derecho > Actualizar campos."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for e in (fld, instr, sep, txt, end):
        run._r.append(e)


def page_number_header():
    hdr = doc.sections[0].header
    par = hdr.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par.paragraph_format.first_line_indent = Cm(0)
    run = par.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    for e in (f1, it, f2):
        run._r.append(e)


page_number_header()

# ==================================================================== PORTADA
for _ in range(4):
    p("", space_after=0)
p("Instalación, configuración y uso de Docker en GNU/Linux",
  align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
p("Adaptación de la guía del SENA a un entorno Ubuntu 24.04 LTS "
  "y aplicación al proyecto CodeByMike",
  align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
p("", space_after=0)
p("Michael David Rodriguez Beltran", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p("Servicio Nacional de Aprendizaje (SENA), Centro de Servicios Financieros, "
  "Regional Distrito Capital", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p("Análisis y Desarrollo de Software, Ficha 3114731", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p("Instructor(a): [NOMBRE DE LA DOCENTE]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p("[FECHA DE ENTREGA]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
doc.add_page_break()

# ============================================================ TABLA CONTENIDO
p("Tabla de Contenido", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
toc()
doc.add_page_break()

# =================================================================== RESUMEN
p("Resumen", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
p("El presente informe documenta el desarrollo del taller de instalación y uso de "
  "Docker propuesto por el SENA, ejecutado sobre un entorno GNU/Linux (Ubuntu 24.04.4 "
  "LTS) en lugar del entorno Windows con Docker Desktop y WSL 2 descrito en la guía. "
  "Se justifica técnicamente cada desviación del procedimiento original, se establece "
  "una tabla de equivalencias entre la interfaz gráfica de Docker Desktop y la interfaz "
  "de línea de comandos de Docker Engine, y se resuelven los seis ejercicios prácticos "
  "de la sección 16 de la guía utilizando la infraestructura de contenedores real del "
  "proyecto CodeByMike (codebymike.tech) en sustitución de la aplicación de ejemplo "
  "Dino Run. Los conceptos evaluados (imagen, contenedor, Dockerfile, contexto de "
  "construcción, publicación de puertos, variables de entorno, bind mount, persistencia, "
  "registros y Docker Compose) se demuestran integros; únicamente cambia la aplicación "
  "sobre la cual se ejercitan y la vía de acceso al motor.", indent=False)
p("", space_after=0)
par = p("", indent=False)
rich(par, [("Palabras clave: ", {"italic": True}),
           ("Docker, contenedores, GNU/Linux, Docker Compose, bind mount, "
            "reproducibilidad, integración continua", {})])
doc.add_page_break()

# ============================================================== INTRODUCCION
doc.add_heading("Introducción", level=1)
p("Docker es una plataforma de virtualización a nivel de sistema operativo que permite "
  "empaquetar una aplicación junto con sus dependencias, su configuración y su comando "
  "de arranque en una unidad denominada imagen, a partir de la cual se ejecutan procesos "
  "aislados llamados contenedores (Docker Inc., 2026a). A diferencia de una máquina "
  "virtual, un contenedor no incorpora un sistema operativo completo: comparte el núcleo "
  "del sistema anfitrión y ejecuta únicamente los procesos de la aplicación, lo que "
  "reduce el consumo de recursos y acelera el arranque.")
p("La guía de formación Instalación y uso de Docker desarrolla este contenido sobre un "
  "entorno Windows, apoyandose en Docker Desktop, en el Subsistema de Windows para Linux "
  "en su versión 2 (WSL 2) y en el asistente Gordon. El equipo de trabajo del autor opera "
  "bajo Ubuntu 24.04.4 LTS, razón por la cual una parte sustancial del procedimiento de "
  "instalación descrito no resulta aplicable: no por dificultad técnica, sino porque el "
  "problema que dicho procedimiento resuelve no se presenta en este entorno.")
p("Este informe no omite contenido de la guía. Documenta, para cada paso del "
  "procedimiento original, su equivalente en GNU/Linux o la razón técnica por la cual no "
  "aplica, y desarrolla los seis ejercicios prácticos propuestos en la sección 16 sobre "
  "la infraestructura de contenedores del proyecto CodeByMike, un portafolio profesional "
  "con panel de control y portal de clientes construido con Astro, Turso/libSQL y "
  "desplegado en Vercel.")
doc.add_page_break()

# ======================================================== 1. JUSTIFICACION
doc.add_heading("Justificación de la Adaptación al Entorno GNU/Linux", level=1)
p("Docker se concibió originalmente como una herramienta de GNU/Linux y su motor se "
  "apoya en dos mecanismos propios del núcleo de ese sistema: los espacios de nombres "
  "(namespaces), que aislan la vista que un proceso tiene del sistema, y los grupos de "
  "control (cgroups), que limitan los recursos que puede consumir (Docker Inc., 2026a). "
  "En consecuencia, para ejecutar contenedores Linux se requiere un núcleo Linux.")
p("WSL 2 es precisamente el mecanismo mediante el cual Windows obtiene ese núcleo: una "
  "máquina virtual ligera y gestionada que ejecuta un núcleo Linux real, sobre la cual "
  "Docker Desktop arranca el motor (Docker Inc., 2026b). Este es el punto central de la "
  "adaptación y conviene enunciarlo con precisión: en el procedimiento descrito por la "
  "guía, la aplicación tampoco se ejecuta sobre Windows, sino sobre una distribución de "
  "Linux alojada en WSL 2. Docker Desktop constituye la capa de integración que oculta "
  "ese hecho al usuario.")
p("Al trabajar directamente sobre Ubuntu, el núcleo requerido es el del propio sistema. "
  "Desaparecen por tanto la instalación de WSL 2, la habilitación de la virtualización "
  "por hardware en BIOS/UEFI y la capa de virtualización intermedia. El motor se instala "
  "como un servicio del sistema administrado por systemd y se opera mediante la interfaz "
  "de línea de comandos. La Tabla 1 detalla la correspondencia completa.")

table(
    "Correspondencia entre el procedimiento de la guía (Windows) y el entorno GNU/Linux",
    ["Paso de la guía", "Situación en GNU/Linux", "Evidencia sustitutiva"],
    [
        ["Instalar Docker Desktop (sección 4)",
         "No aplica. En Linux, Docker Desktop es opcional y solo agrega una interfaz gráfica y una máquina virtual intermedia. Se emplea Docker Engine nativo.",
         "docker --version, docker info con el motor nativo"],
        ["Activar WSL 2 (sección 5.1)",
         "No existe. El núcleo Linux es el del propio sistema operativo.",
         "uname -r, que identifica el núcleo compartido con los contenedores"],
        ["Habilitar virtualización en BIOS/UEFI (sección 5.2)",
         "No se requiere. Los contenedores son procesos aislados mediante namespaces y cgroups, no máquinas virtuales.",
         "docker info mostrando apparmor, seccomp y cgroupns"],
        ["Panel con estado Engine running (sección 5.3)",
         "El motor es un servicio de systemd; el estado se consulta con el gestor de servicios.",
         "systemctl is-active docker y systemctl is-enabled docker"],
        ["Vistas gráficas Builds, Images, Containers, Logs y Bind mounts",
         "No hay interfaz gráfica, pero cada vista tiene un comando equivalente que entrega igual o mayor detalle.",
         "Tabla 3 de equivalencias entre interfaz gráfica y línea de comandos"],
        ["Asistente Gordon (secciones 8.1 y siguientes)",
         "No disponible sin Docker Desktop. La propia guía documenta los comandos que Gordon ejecuta internamente.",
         "docker compose up -d --build y equivalentes, escritos de forma explicita"],
        ["Archivo iniciar_juego.bat",
         "Archivo por lotes propio de Windows; no aplica.",
         "Scripts npm run db:up, db:seed y db:reset definidos en package.json"],
    ],
    note=[("Elaboración propia a partir de la guía ", {}),
          ("Instalación y uso de Docker", {"italic": True}),
          (" (SENA, 2026) y de la documentación oficial de Docker.", {})],
    widths=[4.6, 6.2, 5.2],
)

doc.add_heading("Entorno de Trabajo Verificado", level=2)
p("La totalidad de los comandos y resultados consignados en este informe se ejecutaron y "
  "verificaron en el entorno descrito en la Tabla 2.")

table(
    "Especificación del entorno de trabajo",
    ["Componente", "Versión o estado verificado"],
    [
        ["Sistema operativo", "Ubuntu 24.04.4 LTS"],
        ["Núcleo (kernel)", "7.0.0-28-generic"],
        ["Docker Engine", "29.6.2, build dfc4efb"],
        ["Docker Compose", "v5.3.1 (complemento CLI, no docker-compose de Python)"],
        ["Estado del servicio", "active (en ejecución) y enabled (arranque automático)"],
        ["Acceso sin privilegios", "Usuario perteneciente al grupo docker"],
        ["Opciones de seguridad", "apparmor, seccomp con perfil incorporado, cgroupns"],
        ["Proyecto de aplicación", "CodeByMike: Astro 7 (SSR), Turso/libSQL, Drizzle, Tailwind 4, desplegado en Vercel"],
    ],
    note=[("Datos obtenidos mediante ", {}), ("docker info", {"mono": True}),
          (", ", {}), ("systemctl", {"mono": True}), (" y ", {}),
          ("uname -r", {"mono": True}), (". Elaboración propia.", {})],
    widths=[5.0, 11.0],
)
doc.add_page_break()

# ============================================ 2. INSTALACION Y CONFIGURACION
doc.add_heading("Instalación y Configuración de Docker en Ubuntu", level=1)
p("Esta sección sustituye a las secciones 4, 5 y 15.1 de la guía. Comprende los pasos "
  "que si corresponden a un entorno GNU/Linux, en el orden en que deben ejecutarse.")

doc.add_heading("Instalación desde el Repositorio Oficial", level=2)
p("Los paquetes de Docker incluidos en los repositorios de la distribución suelen "
  "corresponder a versiones antiguas y no incorporan el complemento de Compose en su "
  "versión 2. Por ello, la instalación se realiza desde el repositorio oficial de Docker, "
  "conforme a la documentación del fabricante (Docker Inc., 2026c).")
code([
    "# 1. Retirar paquetes antiguos provenientes de la distribución",
    "sudo apt remove docker docker-engine docker.io containerd runc",
    "",
    "# 2. Registrar la clave GPG y el repositorio oficial",
    "sudo apt update && sudo apt install -y ca-certificates curl",
    "sudo install -m 0755 -d /etc/apt/keyrings",
    "sudo curl -fsSL https://download.docker.com/linux/ubuntu/gpg \\",
    "     -o /etc/apt/keyrings/docker.asc",
    "sudo chmod a+r /etc/apt/keyrings/docker.asc",
    'echo "deb [arch=$(dpkg --print-architecture) \\',
    'signed-by=/etc/apt/keyrings/docker.asc] \\',
    'https://download.docker.com/linux/ubuntu \\',
    '$(. /etc/os-release && echo $VERSION_CODENAME) stable" \\',
    "  | sudo tee /etc/apt/sources.list.d/docker.list > /dev/null",
    "",
    "# 3. Instalar motor, CLI, containerd y los complementos buildx y compose",
    "sudo apt update",
    "sudo apt install -y docker-ce docker-ce-cli containerd.io \\",
    "     docker-buildx-plugin docker-compose-plugin",
])
p("La verificación de la clave GPG no es un formalismo administrativo: garantiza que los "
  "paquetes descargados fueron firmados por Docker y no sustituidos en transito, lo cual "
  "constituye una medida elemental de seguridad de la cadena de suministro.")

doc.add_heading("Configuración Posterior a la Instalación", level=2)
p("Dos ajustes son específicos de GNU/Linux y constituyen el equivalente funcional de "
  "esperar a que Docker Desktop informe el estado Running.")
par = p("", )
rich(par, [("Habilitación del servicio. ", {"bold": True}),
           ("El motor se ejecuta como un demonio administrado por systemd. Debe "
            "distinguirse entre dos estados que la guía trata como uno solo: ", {}),
           ("active", {"mono": True}),
           (" indica que el motor está en ejecución en este momento, mientras que ", {}),
           ("enabled", {"mono": True}),
           (" indica que volverá a iniciarse automáticamente tras reiniciar el equipo. "
            "La evidencia debe mostrar ambos.", {})])
code([
    "sudo systemctl enable --now docker",
    "systemctl is-active docker      # -> active",
    "systemctl is-enabled docker     # -> enabled",
])
par = p("")
rich(par, [("Acceso sin privilegios de superusuario. ", {"bold": True}),
           ("El demonio expone un socket en ", {}),
           ("/var/run/docker.sock", {"mono": True}),
           (", cuyo grupo propietario es ", {}), ("docker", {"mono": True}),
           (". Sin pertenecer a dicho grupo, cada invocación del cliente exigiría ", {}),
           ("sudo", {"mono": True}), (".", {})])
code([
    "sudo usermod -aG docker $USER",
    "newgrp docker            # o cerrar sesión y volver a iniciarla",
    "id -nG | tr ' ' '\\n' | grep -x docker    # -> docker",
])
p("Resulta pertinente advertir, en términos de seguridad, que la pertenencia al grupo "
  "docker equivale en la práctica a disponer de privilegios de superusuario, dado que "
  "cualquier usuario capaz de comunicarse con el socket puede montar el sistema de "
  "archivos raíz dentro de un contenedor privilegiado. No se trata de un error de "
  "configuración, sino de una consecuencia del modelo de arquitectura del demonio. La "
  "alternativa endurecida es el modo sin raíz (rootless), que ejecuta el motor con el "
  "identificador del usuario. Este matiz queda encapsulado tras Docker Desktop en el "
  "procedimiento original y solo se hace visible al operar el motor de forma directa.")

figura(
    "Versiones instaladas y estado del servicio Docker",
    "Terminal (GNOME Terminal), en cualquier directorio.",
    ["docker --version",
     "docker compose version",
     "systemctl is-active docker",
     "systemctl is-enabled docker",
     "uname -r"],
    "Las cinco salidas en una sola terminal: versión 29.6.2 del motor, v5.3.1 de "
    "Compose, active, enabled y el núcleo 7.0.0-28-generic. Esta figura sustituye a "
    "la captura del panel de Docker Desktop con el estado Engine running.",
)

figura(
    "Acceso al motor sin privilegios de superusuario",
    "Terminal, en cualquier directorio.",
    ["id -nG | tr ' ' '\\n' | grep -x docker",
     "docker ps"],
    "La palabra docker como resultado de la primera instrucción y, a continuación, "
    "la tabla de contenedores obtenida sin anteponer sudo. Demuestra que la "
    "configuración posterior a la instalación se completo.",
    altura_cm=5,
)

doc.add_heading("Verificación del Motor", level=2)
p("Los cuatro comandos de verificación propuestos en la sección 5.3 de la guía operan de "
  "manera idéntica en GNU/Linux, ya que corresponden al cliente y no a la interfaz "
  "gráfica.")
code([
    "docker --version",
    "docker compose version",
    "docker info",
    "docker run hello-world",
])
p("A ellos se suman comprobaciones propias del entorno nativo, que en Windows carecerían "
  "de sentido por estar mediadas por la máquina virtual de WSL 2:")
code([
    "uname -r",
    "systemctl status docker --no-pager | head -12",
    "docker info --format 'Servidor: {{.ServerVersion}} | Raíz: {{.DockerRootDir}}'",
    "docker info --format 'Seguridad: {{.SecurityOptions}}'",
])

figura(
    "Salida del comando docker info",
    "Terminal, en cualquier directorio.",
    ["docker info | head -40"],
    "El bloque Server con la versión del motor, el controlador de almacenamiento, el "
    "directorio raíz de datos y, en Security Options, los valores apparmor, seccomp y "
    "cgroupns. Estos últimos evidencian los mecanismos de aislamiento del núcleo y "
    "sustituyen a la captura de virtualización habilitada en BIOS/UEFI de la sección 5.2.",
    altura_cm=8,
)

doc.add_heading("Equivalencias entre la Interfaz Gráfica y la Línea de Comandos", level=2)
p("La Tabla 3 constituye el núcleo metodológico de esta adaptación: acredita que ninguna "
  "de las verificaciones solicitadas por la guía se omitio, sino que se practicó por una "
  "vía distinta. En varios casos, la salida en línea de comandos entrega mayor detalle "
  "que la pantalla correspondiente.")

table(
    "Equivalencia entre las vistas de Docker Desktop y los comandos de Docker Engine",
    ["Vista en Docker Desktop", "Comando equivalente en GNU/Linux"],
    [
        ["Panel principal, estado Engine running", [("systemctl is-active docker", {"mono": True})]],
        ["Images", [("docker images", {"mono": True})]],
        ["Containers", [("docker ps -a", {"mono": True})]],
        ["Builds", [("docker buildx history ls", {"mono": True})]],
        ["Builds, pestaña Logs", [("docker buildx history logs <ID>", {"mono": True})]],
        ["Builds, pestaña Source", [("docker buildx history inspect <ID>", {"mono": True})]],
        ["Containers, pestaña Logs", [("docker logs <nombre>", {"mono": True}), ("  /  ", {}), ("docker compose logs -f", {"mono": True})]],
        ["Containers, pestaña Inspect", [("docker inspect <nombre>", {"mono": True})]],
        ["Containers, pestaña Files", [("docker exec -it <nombre> sh", {"mono": True})]],
        ["Containers, pestaña Stats", [("docker stats", {"mono": True})]],
        ["Containers, pestaña Bind mounts", [("docker inspect --format '{{json .Mounts}}' <nombre>", {"mono": True})]],
        ["Volumes", [("docker volume ls", {"mono": True}), ("  /  ", {}), ("docker volume inspect", {"mono": True})]],
        ["Capas de una imagen", [("docker history <imagen>", {"mono": True})]],
        ["Agrupación de un proyecto de Compose", [("docker compose ps", {"mono": True})]],
    ],
    note=[("Elaboración propia. Los subcomandos ", {}), ("buildx history", {"mono": True}),
          (" fueron verificados en Docker Engine 29.6.2.", {})],
    widths=[7.0, 9.0],
)
doc.add_page_break()

# ================================================ 3. SUSTITUCION DE LA APP
doc.add_heading("Sustitución de la Aplicación de Ejemplo", level=1)
p("La guía emplea Dino Run, una aplicación desarrollada con Flask y SQLite, por cuanto "
  "requiere algún artefacto que contenerizar. El proyecto CodeByMike ya utiliza Docker en "
  "producción de su ciclo de desarrollo, y lo hace con una decisión de arquitectura "
  "explícitamente documentada en el archivo compose.yaml del repositorio:")
q = p("Docker no es el runtime de producción de este proyecto y no debe llegar a serlo. "
      "El sitio lo construye y ejecuta Vercel (SSR sobre Fluid Compute); contenerizar la "
      "aplicación Astro para desplegarla perdería edge, despliegues de vista previa por "
      "solicitud de incorporación de cambios y la reversión automática de ci.yml, a "
      "cambió de nada.", indent=False)
q.paragraph_format.left_indent = Cm(1.27)
p("Delimitar el ámbito de aplicación de una herramienta forma parte del criterio "
  "profesional que la formación busca desarrollar. Lo que si se encuentra contenerizado, "
  "y constituye el uso serio de la tecnología en este proyecto, comprende dos artefactos:")
bullets([
    [("compose.yaml", {"mono": True}),
     (" declara dos servidores libSQL (sqld), el mismo motor de base de datos que opera "
      "en producción a través de Turso. Permite ejercitar transacciones, restricciones "
      "UNIQUE y concurrencia contra el protocolo real y no contra un archivo local.", {})],
    [(".devcontainer/Dockerfile", {"mono": True}),
     (" fija el entorno de desarrollo con Node 22.12 exacto y el navegador Chromium de "
      "Playwright, resolviendo un defecto real y documentado: la presencia de Node 20 en "
      "la ruta de ejecución interrumpe la compilación con astro build.", {})],
])
p("Los seis ejercicios se desarrollan sobre estos artefactos. La correspondencia "
  "conceptual es directa, según se detalla en la Tabla 4.")

table(
    "Correspondencia conceptual entre la aplicación de la guía y el proyecto CodeByMike",
    ["Concepto", "En Dino Run (guía)", "En CodeByMike"],
    [
        ["Imagen", "sena-docker-flask-app-app:latest",
         "ghcr.io/tursodatabase/libsql-server, fijada por digest, y la imagen del entorno de desarrollo"],
        ["Contenedor", "dino-game", "codebymike-libsql-main y codebymike-libsql-demo"],
        ["Dockerfile", "Python 3.12, Flask", ".devcontainer/Dockerfile, Node 22.12 y Chromium"],
        ["Puerto publicado", "5000:5000", "127.0.0.1:8080:8080 y 127.0.0.1:8081:8080"],
        ["Persistencia", "bind mount ./data:/app/data (SQLite)",
         "volumen gestionado libsql-main-data:/var/lib/sqld, y bind mount en el ejercicio 5"],
        ["Variables de entorno", "DATABASE_PATH, SECRET_KEY",
         "SQLD_NODE y SQLD_DB_PATH en el contenedor; TURSO_DATABASE_URL y TURSO_DEMO_URL en la aplicación"],
        ["Docker Compose", "Un servicio denominado app",
         "Dos servicios que comparten configuración mediante un ancla YAML (x-libsql)"],
        ["Comprobación de salud", "HEALTHCHECK sobre la página principal",
         "Espera activa en scripts/wait-libsql.mjs hasta que el extremo /health responda"],
    ],
    note=[("Elaboración propia. Todos los conceptos evaluados por la guía se conservan; "
           "únicamente cambia la aplicación sobre la cual se ejercitan.", {})],
    widths=[3.2, 5.4, 7.4],
)
doc.add_page_break()

# ================================================== 4. EJERCICIOS PRACTICOS
doc.add_heading("Desarrollo de los Ejercicios Prácticos", level=1)
p("Los seis ejercicios corresponden a la sección 16 de la guía (página 27). Todos los "
  "comandos se ejecutan desde la raíz del repositorio del proyecto.")

# ---- Ejercicio 1
doc.add_heading("Ejercicio 1. Verificación Inicial del Motor", level=2)
par = p("")
rich(par, [("Problema planteado. ", {"italic": True}),
           ("Comprobar que el motor funciona y puede ejecutar imagenes.", {})])
p("En ausencia de interfaz gráfica, el ciclo completo consiste en descargar una imagen, "
  "crear un contenedor a partir de ella y ejecutarlo.")
code(["docker run hello-world",
      "docker images | head",
      "docker ps -a | head"])
p("Conviene precisar que docker run no constituye una operación única: localiza la imagen "
  "en el almacen local, la descarga desde el registro si no se encuentra presente, crea "
  "el contenedor y lo inicia. La imagen hello-world evidencia las cuatro operaciones de "
  "manera simultánea y por ello la guía la emplea como prueba de funcionamiento.")
p("Dado que dicho contenedor finaliza de inmediato, se complementa la evidencia con un "
  "servicio persistente que publique un puerto, lo que permite verificar además la "
  "conectividad de red:")
code(["docker run -d --name prueba-sena -p 8085:80 nginx",
      "curl -s -o /dev/null -w 'HTTP %{http_code}\\n' http://localhost:8085",
      "docker ps --filter name=prueba-sena",
      "docker logs prueba-sena",
      "docker rm -f prueba-sena"])

figura(
    "Ejecución de la imagen de prueba hello-world",
    "Terminal, en cualquier directorio.",
    ["docker run hello-world"],
    "El mensaje Hello from Docker! junto con la explicación de los cuatro pasos que "
    "el motor ejecutó. Si la imagen no estaba descargada previamente, deben apreciarse "
    "también las líneas de descarga (Unable to find image ... Pulling from ...), lo cual "
    "enriquece la evidencia.",
    altura_cm=8,
)

figura(
    "Contenedor con puerto publicado y verificación por HTTP",
    "Terminal, en cualquier directorio.",
    ["docker run -d --name prueba-sena -p 8085:80 nginx",
     "curl -s -o /dev/null -w 'HTTP %{http_code}\\n' http://localhost:8085",
     "docker ps --filter name=prueba-sena"],
    "El identificador del contenedor creado, la respuesta HTTP 200 y la fila del "
    "contenedor en ejecución con la asignación de puertos 0.0.0.0:8085->80/tcp. "
    "Equivale a pulsar el enlace del puerto en la vista Containers de Docker Desktop.",
    altura_cm=6,
)

# ---- Ejercicio 2
doc.add_heading("Ejercicio 2. Construcción y Revisión de la Imagen", level=2)
par = p("")
rich(par, [("Problema planteado. ", {"italic": True}),
           ("Construir la aplicación y comprender el resultado desde las vistas Images y "
            "Builds.", {})])
p("Se construye el Dockerfile del entorno de desarrollo del proyecto, que constituye un "
  "artefacto en uso y no un ejemplo didáctico:")
code(["docker build -t codebymike-dev:sena .devcontainer/"])
p("La revisión posterior cubre exactamente la información que la interfaz gráfica "
  "presenta en las vistas Builds e Images:")
code(["# Vista Builds: historial de construcciones, estado y duración",
      "docker buildx history ls",
      "",
      "# Vista Builds, pestañas Logs y Source",
      "docker buildx history logs <BUILD_ID>",
      "docker buildx history inspect <BUILD_ID>",
      "",
      "# Vista Images",
      "docker images codebymike-dev",
      "",
      "# Capas de la imagen, con mayor detalle que la interfaz gráfica",
      "docker history codebymike-dev:sena --format 'table {{.Size}}\\t{{.CreatedBy}}'"])
p("La demostración de la memoria caché de capas resulta especialmente pertinente, por "
  "cuanto la sección 6.2 de la guía justifica precisamente con ese argumento el orden de "
  "las instrucciones (copiar el archivo de dependencias antes que el código fuente). Al "
  "repetir la misma construcción, todos los pasos deben resolverse desde la caché y el "
  "proceso concluye en segundos:")
code(["docker build -t codebymike-dev:sena .devcontainer/   # todos los pasos CACHED"])
p("El análisis instrucción por instrucción del Dockerfile empleado, contrastado con el de "
  "Dino Run, se presenta en la Tabla 5.")

table(
    "Análisis del Dockerfile del proyecto y su correspondencia con el de la guía",
    ["Instrucción", "Función en el proyecto", "Correspondencia con la guía"],
    [
        ["FROM node:22.12-bookworm@sha256:0e910f...",
         "Imagen base fijada por digest criptográfico y no por etiqueta.",
         "Equivale a FROM python:3.12-slim, con un criterio más estricto: una etiqueta es mutable, un digest no. Constituye reproducibilidad verificable."],
        ["ARG PLAYWRIGHT_VERSION=1.61.1",
         "Fija la versión del navegador de pruebas automatizadas.",
         "Sin equivalente en la guía. Evita que Playwright rechace un navegador instalado por otra versión."],
        ["RUN npx playwright install-deps chromium",
         "Instala dependencias de sistema con privilegios de root.",
         "Equivale a RUN pip install --no-cache-dir -r requirements.txt."],
        ["USER node",
         "Reduce privilegios antes de ejecutar el proceso principal.",
         "Idéntico a USER appuser, con idéntico fundamento: limitar el impacto de una eventual vulnerabilidad."],
        ["ENV PLAYWRIGHT_BROWSERS_PATH=...",
         "Variable de entorno incorporada a la imagen.",
         "La guía prefiere declarar variables en compose.yaml. Esta pertenece a la imagen porque no varía entre entornos."],
        ["WORKDIR /workspace",
         "Establece el directorio de trabajo dentro del contenedor.",
         "Equivale a WORKDIR /app."],
    ],
    note=[("Elaboración propia a partir de ", {}), (".devcontainer/Dockerfile", {"mono": True}),
          (" del repositorio del proyecto.", {})],
    widths=[4.4, 5.4, 6.2],
)

p("En cuanto al contexto de construcción, el archivo .dockerignore del proyecto responde "
  "a la pregunta formulada en la sección 6.3 de la guía con un argumento que allí se "
  "menciona de forma tangencial y que en el repositorio consta como comentario del propio "
  "archivo: además del peso del contexto, se trata de higiene de la cadena de suministro, "
  "puesto que un secreto que ingresa al contexto puede permanecer en una capa de la imagen "
  "aunque el Dockerfile no lo copie en ningún momento. Se excluyen .env y sus variantes, "
  ".vercel, .git, node_modules, dist, .astro, los informes de cobertura y de pruebas, y "
  "los documentos en formato PDF y DOCX.")

figura(
    "Primera construcción de la imagen del entorno de desarrollo",
    "Terminal, en la raíz del repositorio del proyecto.",
    ["docker build -t codebymike-dev:sena .devcontainer/"],
    "La secuencia completa de pasos numerados, la transferencia del contexto de "
    "construcción, la descarga de la imagen base y la línea final con el tiempo total. "
    "Equivale a la vista Builds con estado Completed.",
    altura_cm=9,
)

figura(
    "Segunda construcción con aprovechamiento de la memoria caché",
    "Terminal, en la raíz del repositorio, inmediatamente después de la Figura 6.",
    ["docker build -t codebymike-dev:sena .devcontainer/"],
    "La totalidad de los pasos marcados como CACHED y una duración total inferior a "
    "un segundo. Esta figura demuestra el concepto de capas reutilizables que la guía "
    "expone en la sección 6.2.",
    altura_cm=7,
)

figura(
    "Historial de construcciones, imagenes y capas",
    "Terminal, en la raíz del repositorio.",
    ["docker buildx history ls",
     "docker images codebymike-dev",
     "docker history codebymike-dev:sena --format 'table {{.Size}}\\t{{.CreatedBy}}'"],
    "El registro de construcción con estado Completed y su duración (vista Builds), la "
    "fila de la imagen con etiqueta, identificador y tamaño (vista Images), y el desglose "
    "de capas con el tamaño que cada instrucción aporta al total.",
    altura_cm=9,
)

# ---- Ejercicio 3
doc.add_heading("Ejercicio 3. Ejecución y Comprobación con Docker Compose", level=2)
par = p("")
rich(par, [("Problema planteado. ", {"italic": True}),
           ("Ejecutar la aplicación mediante Docker Compose y acceder a ella sin "
            "configurar manualmente cada opción de docker run.", {})])
p("El proyecto encapsula esta operación en scripts de npm, que constituyen la versión "
  "profesional del archivo iniciar_juego.bat propuesto por la guía:")
code(["npm run db:up      # docker compose up -d y espera activa hasta que /health responda",
      "docker compose ps"])
p("De manera equivalente, invocando directamente la herramienta:")
code(["docker compose up -d",
      "docker compose ps --format 'table {{.Name}}\\t{{.Status}}\\t{{.Ports}}'"])
p("La comprobación a través del puerto publicado equivale a pulsar el enlace 5000:5000 "
  "en la vista Containers:")
code(["curl -s -o /dev/null -w 'principal HTTP %{http_code}\\n' http://127.0.0.1:8080/health",
      "curl -s -o /dev/null -w 'demo HTTP %{http_code}\\n' http://127.0.0.1:8081/health"])
p("Y la aplicación completa operando contra dichas bases de datos:")
code(["npm run db:seed",
      "npm run dev:carga    # Astro contra sqld en los puertos 8080 y 8081"])
p("El archivo compose.yaml del proyecto excede en varios aspectos al presentado por la "
  "guía, y conviene destacarlo:")
bullets([
    [("La construcción ", {}), ("x-libsql: &libsql", {"mono": True}),
     (" junto con ", {}), ("<<: *libsql", {"mono": True}),
     (" corresponde a un ancla y una fusión de YAML: el bloque de configuración común se "
      "declara una sola vez y ambos servicios lo heredan. La guía, al contar con un único "
      "servicio, no llega a plantear este problema.", {})],
    [("La declaración ", {}), ("ports: '127.0.0.1:8080:8080'", {"mono": True}),
     (" publica el servicio exclusivamente en la interfaz de bucle local. La guía escribe ", {}),
     ('"5000:5000"', {"mono": True}),
     (", forma que en GNU/Linux expone el servicio en todas las interfaces de red y lo "
      "deja accesible desde la red local. Se trata de una diferencia de seguridad "
      "sustantiva y no de una preferencia de estilo.", {})],
    [("Las directivas ", {}), ("cap_drop: ALL", {"mono": True}), (" y ", {}),
     ("cap_add: [DAC_OVERRIDE, CHOWN, SETUID, SETGID]", {"mono": True}),
     (" restringen el contenedor al conjunto mínimo de capacidades del núcleo que sqld "
      "requiere para iniciar. El archivo documenta el método empleado para determinarlas, "
      "consistente en retirarlas una a una hasta que el extremo /health dejó de responder.", {})],
    [("La directiva ", {}), ("restart: unless-stopped", {"mono": True}),
     (" es idéntica a la propuesta por la guía.", {})],
])

figura(
    "Servicios de Compose en ejecución y respuesta de ambas bases de datos",
    "Terminal, en la raíz del repositorio.",
    ["docker compose ps --format 'table {{.Name}}\\t{{.Status}}\\t{{.Ports}}'",
     "curl -s -o /dev/null -w 'principal HTTP %{http_code}\\n' http://127.0.0.1:8080/health",
     "curl -s -o /dev/null -w 'demo HTTP %{http_code}\\n' http://127.0.0.1:8081/health"],
    "Los dos contenedores con estado Up y sus puertos publicados en 127.0.0.1, seguidos "
    "de las dos respuestas HTTP 200. Equivale a la vista Containers con el grupo de "
    "Compose desplegado y ambos servicios en estado Running.",
    altura_cm=6,
)

figura(
    "Aplicación web en ejecución contra las bases de datos en contenedores",
    "Navegador web, dirección http://localhost:4400, tras ejecutar npm run dev:carga.",
    ["npm run db:seed",
     "npm run dev:carga"],
    "Una página del proyecto que muestre datos provenientes de la base de datos (por "
    "ejemplo el panel de control o la página de estado), lo cual acredita que la "
    "aplicación consume efectivamente los servicios contenerizados. Equivale a la "
    "captura de Dino Run en el navegador.",
    altura_cm=9,
)

# ---- Ejercicio 4
doc.add_heading("Ejercicio 4. Variables de Entorno y Gestión de Secretos", level=2)
par = p("")
rich(par, [("Problema planteado. ", {"italic": True}),
           ("Distinguir las variables predeterminadas, opcionales y sensibles, y "
            "definirlas correctamente para Docker Compose, comprobando únicamente que la "
            "variable existe, sin mostrar ni capturar su valor real.", {})])
p("Las variables declaradas en compose.yaml y entregadas al contenedor se consultan así:")
code(["docker inspect codebymike-libsql-main \\",
      "  --format '{{range .Config.Env}}{{println .}}{{end}}'",
      "",
      "# SQLD_NODE=primary",
      "# SQLD_DB_PATH=/var/lib/sqld/main.db"])
p("La variable SQLD_DB_PATH constituye el paralelo exacto de DATABASE_PATH en Dino Run: "
  "indica al servicio la ruta en la cual debe escribir, y dicha ruta debe coincidir con "
  "el punto de montaje declarado. Si una de las dos se modifica sin la otra, los datos se "
  "escriben fuera del volumen y se pierden al recrear el contenedor.")
p("Del lado de la aplicación, el repositorio establece un criterio propio, consignado en "
  "su archivo de instrucciones técnicas: dado que el proyecto combina import.meta.env y "
  "process.env, que no son equivalentes (el servidor de desarrollo carga el archivo .env "
  "solo en el primero, mientras que Vercel inyecta las variables únicamente en el "
  "segundo), toda lectura de una variable nueva debe realizarse mediante la función "
  "serverEnv de src/lib/env.ts, que consulta ambas fuentes. La clasificación solicitada "
  "por el ejercicio se presenta en la Tabla 6.")

table(
    "Clasificación de las variables de entorno del proyecto",
    ["Tipo", "Ejemplos", "Tratamiento"],
    [
        ["Predeterminada",
         "TURSO_DATABASE_URL=http://127.0.0.1:8080",
         "Declarada en el script dev:carga de package.json. Su valor es conocido, no comprometedor y puede publicarse."],
        ["Opcional",
         "NTFY_TOPIC, RESEND_API_KEY",
         "Si la variable no existe, el módulo src/lib/notify.ts retorna un resultado con la marca skipped y nunca interrumpe la ejecución. Equivale a la variable FLASK_ENV que la guía califica de opcional, con la diferencia de que aquí la degradación está diseñada de forma deliberada."],
        ["Sensible",
         "ENCRYPTION_KEY, AUTH_SECRET, CRON_SECRET, COBRO_HISTORY_SECRET",
         "Residen en el archivo .env local, excluido por .gitignore y por .dockerignore, y en las variables de entorno de la plataforma de despliegue. Nunca se escriben en compose.yaml ni aparecen en una captura de pantalla."],
    ],
    note=[("Elaboración propia. El criterio de tratamiento corresponde a las reglas de "
           "arquitectura documentadas en el repositorio del proyecto.", {})],
    widths=[2.8, 4.4, 8.8],
)

p("La comprobación solicitada exige acreditar la existencia de la variable sin revelar su "
  "valor. La forma correcta de hacerlo en una evidencia que se entrega consiste en "
  "recortar el contenido posterior al signo igual:")
code(["# Correcto: se prueba la existencia, no se expone el valor",
      "grep -c '^ENCRYPTION_KEY=' .env          # -> 1",
      "",
      "docker inspect codebymike-libsql-main \\",
      "  --format '{{range .Config.Env}}{{println .}}{{end}}' | cut -d= -f1",
      "# SQLD_NODE",
      "# SQLD_DB_PATH",
      "# PATH"])
p("Resulta oportuno señalar que la guía incurre en su propio ejemplo en la práctica que "
  "posteriormente advierte: publica el valor SECRET_KEY=tu-clave-secreta-mejor-cambiar-en-"
  "produccion dentro del archivo compose.yaml reproducido en el documento. La propia guía "
  "lo corrige en la sección 12, al recomendar la interpolación desde un archivo .env "
  "local mediante la sintaxis SECRET_KEY=${SECRET_KEY}. El proyecto CodeByMike aplica esa "
  "regla de forma sistemática: ningún secreto figura escrito en compose.yaml.")

figura(
    "Comprobación de variables de entorno sin exposición de valores",
    "Terminal, en la raíz del repositorio.",
    ["docker inspect codebymike-libsql-main \\",
     "  --format '{{range .Config.Env}}{{println .}}{{end}}'",
     "",
     "docker inspect codebymike-libsql-main \\",
     "  --format '{{range .Config.Env}}{{println .}}{{end}}' | cut -d= -f1",
     "",
     "grep -c '^ENCRYPTION_KEY=' .env",
     "grep -nE '^\\.env' .dockerignore .gitignore"],
    "Las variables del contenedor con sus valores (no son secretos: SQLD_NODE y "
    "SQLD_DB_PATH), la misma consulta recortada a solo los nombres, el conteo que "
    "acredita la existencia de ENCRYPTION_KEY sin mostrarla, y las líneas de "
    ".dockerignore y .gitignore que excluyen el archivo .env. Verificar antes de "
    "capturar que no haya ningún valor secreto visible en la terminal.",
    altura_cm=8,
)

# ---- Ejercicio 5
doc.add_heading("Ejercicio 5. Persistencia mediante Bind Mount", level=2)
par = p("")
rich(par, [("Problema planteado. ", {"italic": True}),
           ("Conservar los datos mediante el bind mount declarado en compose.yaml: "
            "registrar información, reiniciar el contenedor y comprobar que permanece.", {})])
p("El proyecto emplea volúmenes gestionados por Docker y no montajes de enlace, decisión "
  "que resulta apropiada para datos de servicio: Docker administra los permisos y el "
  "ciclo de vida del almacenamiento, y los archivos no contaminan el árbol del "
  "repositorio. Para desarrollar el ejercicio se dispuso un archivo de anulación que "
  "modifica únicamente ese aspecto, sin alterar el comportamiento predeterminado del "
  "proyecto. El archivo docs/manuales-sena/compose.sena-bind.yaml contiene:")
code(["services:",
      "  libsql-main:",
      "    volumes:",
      "      - ./.data/libsql-main:/var/lib/sqld"])
p("El procedimiento completo, con el orden de las comprobaciones, es el siguiente:")
code(["# 1. Estado inicial: volumen gestionado por Docker",
      "docker inspect codebymike-libsql-main \\",
      "  --format '{{range .Mounts}}{{.Type}} {{.Source}} -> {{.Destination}}{{println}}{{end}}'",
      "",
      "# 2. Levantar los servicios con el archivo de anulación",
      "docker compose -f compose.yaml \\",
      "  -f docs/manuales-sena/compose.sena-bind.yaml up -d",
      "",
      "# 3. Confirmar el cambio de tipo de montaje (vista Bind mounts)",
      "docker inspect codebymike-libsql-main \\",
      "  --format '{{range .Mounts}}{{.Type}} {{.Source}} -> {{.Destination}}{{println}}{{end}}'",
      "",
      "# 4. La carpeta aparece en el anfitrión, que es el propósito del bind mount",
      "ls -la .data/libsql-main",
      "",
      "# 5. Escribir datos, reiniciar y comprobar",
      "npm run db:seed",
      "docker compose restart libsql-main",
      "du -sh .data/libsql-main"])
p("Para una comprobación más explícita, equivalente a registrar un usuario y verificar "
  "que su puntuación permanece tras el reinicio, se inserta un registro, se reinicia el "
  "contenedor y se consulta nuevamente:")
code(["node -e \"",
      "const {createClient} = require('@libsql/client');",
      "const c = createClient({url:'http://127.0.0.1:8080'});",
      "(async () => {",
      "  await c.execute('CREATE TABLE IF NOT EXISTS prueba_sena(id INTEGER PRIMARY KEY, nota TEXT)');",
      "  await c.execute(\\\"INSERT INTO prueba_sena(nota) VALUES('antes del reinicio')\\\");",
      "  console.log('filas:', (await c.execute('SELECT count(*) n FROM prueba_sena')).rows[0].n);",
      "})();\"",
      "",
      "docker compose restart libsql-main && sleep 5",
      "",
      "node -e \"",
      "const {createClient} = require('@libsql/client');",
      "createClient({url:'http://127.0.0.1:8080'})",
      "  .execute('SELECT count(*) n, max(nota) nota FROM prueba_sena')",
      "  .then(r => console.log('tras el reinicio:', r.rows[0].n, '|', r.rows[0].nota));\""])
p("El resultado obtenido en el entorno descrito fue el siguiente:")
code(["filas: 1",
      "tras el reinicio: 1 | antes del reinicio"])
p("Para restablecer el modo predeterminado del proyecto se recrean los servicios sin el "
  "archivo de anulación:")
code(["docker compose up -d --force-recreate"])
par = p("")
rich(par, [("Hallazgo específico del entorno GNU/Linux. ", {"bold": True}),
           ("Los archivos que sqld crea dentro de la carpeta enlazada quedan bajo el "
            "identificador de usuario del contenedor (666) y no bajo el del usuario del "
            "anfitrión. En consecuencia, su eliminación desde el sistema anfitrión falla "
            "con el mensaje Permission denied, y debe realizarse desde un contenedor, "
            "que si dispone de privilegios de superusuario:", {})])
code(['docker run --rm -v "$PWD/.data:/x" alpine rm -rf /x/libsql-main'])
p("Este comportamiento no se manifiesta en Windows con Docker Desktop, dado que el acceso "
  "al sistema de archivos se realiza a través de una capa de traducción que homogeniza la "
  "propiedad de los archivos. Constituye una de las diferencias reales entre ambas "
  "plataformas y explica por que en equipos GNU/Linux se prefieren volúmenes gestionados "
  "para datos de servicio, reservando los montajes de enlace para el código fuente "
  "durante el desarrollo.")

figura(
    "Cambió del tipo de montaje: de volumen gestionado a bind mount",
    "Terminal, en la raíz del repositorio.",
    ["docker inspect codebymike-libsql-main \\",
     "  --format '{{range .Mounts}}{{.Type}} {{.Source}} -> {{.Destination}}{{println}}{{end}}'",
     "",
     "docker compose -f compose.yaml \\",
     "  -f docs/manuales-sena/compose.sena-bind.yaml up -d",
     "",
     "docker inspect codebymike-libsql-main \\",
     "  --format '{{range .Mounts}}{{.Type}} {{.Source}} -> {{.Destination}}{{println}}{{end}}'",
     "",
     "ls -la .data/libsql-main"],
    "Las dos consultas de montaje una debajo de la otra: la primera con el prefijo "
    "volume y una ruta bajo /var/lib/docker/volumes, la segunda con el prefijo bind y la "
    "ruta del repositorio. Al final, el listado de la carpeta creada en el anfitrión. "
    "Equivale a la pestaña Bind mounts de Docker Desktop.",
    altura_cm=8,
)

figura(
    "Comprobación de la persistencia de los datos tras el reinicio del contenedor",
    "Terminal, en la raíz del repositorio.",
    ["# inserción, reinicio y nueva consulta (bloques node -e de esta sección)",
     "docker compose restart libsql-main",
     "du -sh .data/libsql-main"],
    "La secuencia completa en una sola terminal: filas: 1, el reinicio del contenedor y "
    "el resultado tras el reinicio: 1 | antes del reinicio. Esta figura acredita el "
    "criterio de evaluación relativo a la persistencia.",
    altura_cm=7,
)

# ---- Ejercicio 6
doc.add_heading("Ejercicio 6. Diagnóstico de un Puerto Ocupado", level=2)
par = p("")
rich(par, [("Problema planteado. ", {"italic": True}),
           ("Diagnosticar un conflicto de puerto y adaptar el puerto del anfitrión en "
            "compose.yaml.", {})])
p("El error se reproduce de forma deliberada y no es necesario esperar a que ocurra. Con "
  "el contenedor codebymike-libsql-main ocupando el puerto 8080, basta solicitar el mismo "
  "puerto para otro contenedor:")
code(["docker run -d --name choque -p 127.0.0.1:8080:80 nginx"])
p("El mensaje obtenido en el entorno descrito fue:")
code(["Error response from daemon: failed to set up container networking:",
      "driver failed programming external connectivity on endpoint choque:",
      "Bind for 127.0.0.1:8080 failed: port is already allocated"])
p("El diagnóstico, equivalente a identificar en la vista Containers cual servicio ocupa "
  "el puerto, se realiza en dos niveles:")
code(["# Si el puerto lo ocupa un contenedor",
      "docker ps --filter publish=8080 --format '{{.Names}} -> {{.Ports}}'",
      "# codebymike-libsql-main -> 5001/tcp, 127.0.0.1:8080->8080/tcp",
      "",
      "# Si lo ocupa un proceso nativo del anfitrión",
      "ss -ltnp | grep :8080"])
p("El segundo comando constituye una ventaja del entorno GNU/Linux: cuando el puerto se "
  "encuentra ocupado por un proceso del sistema anfitrión y no por un contenedor, Docker "
  "Desktop no puede identificarlo y el diagnóstico queda incompleto.")
p("La solución propuesta por la guía consiste en modificar el puerto del anfitrión y no "
  "el del contenedor:")
code(["docker rm -f choque",
      "docker run -d --name choque -p 127.0.0.1:8086:80 nginx",
      "curl -s -o /dev/null -w 'HTTP %{http_code}\\n' http://127.0.0.1:8086",
      "docker rm -f choque"])
p("En la notación ANFITRIÓN:CONTENEDOR, el valor del lado derecho lo determina la "
  "aplicación (sqld escucha en el puerto 8080 y no admite negociación), mientras que el "
  "del lado izquierdo es de libre elección. Por esa razón la corrección consiste siempre "
  "en desplazar el valor izquierdo, y por esa misma razón la instrucción EXPOSE no "
  "publica puerto alguno: únicamente documenta el puerto interno.")

figura(
    "Conflicto de puerto, diagnóstico y resolución",
    "Terminal, en la raíz del repositorio.",
    ["docker run -d --name choque -p 127.0.0.1:8080:80 nginx",
     "docker ps --filter publish=8080 --format '{{.Names}} -> {{.Ports}}'",
     "ss -ltnp | grep :8080",
     "docker rm -f choque",
     "docker run -d --name choque -p 127.0.0.1:8086:80 nginx",
     "curl -s -o /dev/null -w 'HTTP %{http_code}\\n' http://127.0.0.1:8086",
     "docker rm -f choque"],
    "La secuencia completa: el mensaje port is already allocated, la identificación del "
    "contenedor que ocupa el puerto, la confirmación a nivel de sistema operativo, y el "
    "mismo contenedor operando correctamente en el puerto alterno con respuesta HTTP 200.",
    altura_cm=9,
)
doc.add_page_break()

# ================================================ 5. ACTIVIDADES DE APRENDIZAJE
doc.add_heading("Actividades de Aprendizaje", level=1)

doc.add_heading("Reflexión Inicial", level=2)
p("Empaquetar una aplicación consiste en entregar el programa junto con la totalidad de "
  "los elementos que requiere para ejecutarse (intérprete o máquina de ejecución, "
  "bibliotecas, configuración y comando de arranque) dentro de una unidad que no depende "
  "de la configuración previa de la máquina de destino.")
p("El caso concreto en el cual una aplicación funciona en un equipo y falla en otro se "
  "encuentra documentado en el propio repositorio del proyecto y coincide exactamente con "
  "el problema que la guía enuncia en su introducción: el intérprete de Node presente por "
  "defecto en la ruta de ejecución del sistema corresponde a la versión 20, mientras que "
  "la compilación mediante astro build exige la versión 22.12 o superior. El mismo "
  "repositorio y el mismo comando producen resultados distintos según la máquina. El "
  "directorio .devcontainer constituye la respuesta a ese problema: fija la versión 22.12 "
  "dentro de una imagen y elimina esa clase completa de fallo.")

doc.add_heading("Contextualización y Conocimientos Previos", level=2)
bullets([
    [("Dockerfile. ", {"bold": True}),
     ("Archivo de texto que contiene la receta de construcción. En este proyecto, "
      ".devcontainer/Dockerfile, versionado junto con el código fuente.", {})],
    [("Imagen. ", {"bold": True}),
     ("Resultado de ejecutar esa receta. Es inmutable y está organizada en capas "
      "superpuestas y reutilizables mediante memoria caché. Se consulta con docker images.", {})],
    [("Contenedor. ", {"bold": True}),
     ("Instancia en ejecución de una imagen, con una capa de escritura propia. De una "
      "misma imagen se derivan múltiples contenedores: codebymike-libsql-main y "
      "codebymike-libsql-demo ejecutan la misma imagen con configuración distinta, hecho "
      "que se aprecia directamente en la salida de docker ps.", {})],
])

doc.add_heading("Apropiación del Conocimiento", level=2)
p("Los contenidos de apropiación quedan cubiertos por los ejercicios 2 (capas y memoria "
  "cache), 3 (orquestación con Compose) y 5 (persistencia). La actividad consistente en "
  "modificar un elemento y reconstruir la imagen se traduce, en este proyecto, en alterar "
  "el valor de ARG PLAYWRIGHT_VERSION del Dockerfile del entorno de desarrollo y observar "
  "en la construcción siguiente que capas se invalidan y cuales se conservan en caché: la "
  "modificación invalida esa instrucción y todas las posteriores, mientras que la imagen "
  "base declarada en FROM no vuelve a descargarse.")
doc.add_page_break()

# ==================================================== 6. EVIDENCIAS
doc.add_heading("Relación de Evidencias y Criterios de Evaluación", level=1)
p("La Tabla 7 relaciona cada figura del presente informe con el criterio de evaluación de "
  "la sección 19 de la guía que dicha evidencia acredita.")

table(
    "Relación entre las evidencias gráficas y los criterios de evaluación",
    ["Figura", "Evidencia", "Criterio de evaluación acreditado"],
    [
        ["1 y 2", "Motor instalado, activo, habilitado y accesible sin privilegios de superusuario",
         "Instala Docker y se orienta correctamente en su entorno de operación"],
        ["3", "Salida de docker info con los mecanismos de aislamiento del núcleo",
         "Comprende los componentes del motor y su modelo de aislamiento"],
        ["4 y 5", "Ejecución de hello-world y de un contenedor con puerto publicado",
         "Ejecuta imagenes y verifica el funcionamiento del motor"],
        ["6, 7 y 8", "Construcción de la imagen, aprovechamiento de la caché y revisión de capas",
         "Construye la imagen y analiza el proceso desde el historial de construcciones"],
        ["9 y 10", "Servicios de Compose en ejecución y aplicación accesible por el puerto publicado",
         "Ejecuta y comprueba el proyecto mediante Compose y el puerto publicado"],
        ["11", "Variables de entorno verificadas sin exposición de valores sensibles",
         "Distingue variables opcionales, predeterminadas y sensibles; entrega evidencias sin exponer información sensible"],
        ["12 y 13", "Cambió de tipo de montaje y persistencia comprobada tras el reinicio",
         "Configura y verifica la persistencia mediante el montaje declarado en compose.yaml"],
        ["14", "Conflicto de puerto, diagnóstico y resolución",
         "Interpreta registros y corrige errores frecuentes"],
    ],
    note=[("Elaboración propia a partir de los criterios enunciados en la sección 19 de "
           "la guía del SENA.", {})],
    widths=[1.8, 6.2, 8.0],
)
p("Antes de la entrega debe verificarse que ninguna captura exhiba el valor de una "
  "variable sensible. El propio criterio de evaluación lo exige de forma expresa y, en "
  "este proyecto, constituye además una regla de arquitectura documentada.")
doc.add_page_break()

# ==================================================== 7. CONCLUSIONES
doc.add_heading("Conclusiones", level=1)
p("El desarrollo del taller sobre un entorno GNU/Linux no supuso una reducción del "
  "alcance formativo sino un desplazamiento del punto de observación. Los conceptos "
  "evaluados por la guía (imagen, contenedor, Dockerfile, contexto de construcción, "
  "publicación de puertos, variables de entorno, montajes, persistencia, registros y "
  "orquestación mediante Compose) se ejercitaron en su totalidad; lo que cambió fue la "
  "vía de acceso al motor y la aplicación empleada como caso de estudio.")
p("Operar Docker Engine de forma directa expuso elementos que la interfaz gráfica "
  "encapsula y que resultan pertinentes para la formación profesional: la naturaleza del "
  "motor como servicio del sistema, el modelo de permisos basado en el socket del demonio "
  "y sus implicaciones de seguridad, la propiedad de los archivos generados en montajes "
  "de enlace, y la posibilidad de diagnosticar conflictos de puerto a nivel del sistema "
  "operativo y no únicamente entre contenedores.")
p("La aplicación del taller a un proyecto real permitió además contrastar el ejemplo "
  "didáctico con prácticas de mayor exigencia: la fijación de imagenes por digest "
  "criptográfico frente al uso de etiquetas mutables, la publicación de puertos "
  "restringida a la interfaz de bucle local, la reducción del conjunto de capacidades del "
  "núcleo otorgadas al contenedor, y la exclusión sistemática de secretos del contexto de "
  "construcción. Finalmente, la delimitación del ámbito de aplicación de la herramienta, "
  "es decir, el reconocimiento de que la contenerización resuelve el entorno de "
  "desarrollo y de pruebas de este proyecto pero no su despliegue en producción, "
  "constituye en si misma un resultado de aprendizaje.")
doc.add_page_break()

# ==================================================== REFERENCIAS
doc.add_heading("Referencias", level=1)
refs = [
    "Docker Inc. (2026a). Docker overview. Docker Docs. "
    "https://docs.docker.com/get-started/docker-overview/",
    "Docker Inc. (2026b). Docker Desktop WSL 2 backend on Windows. Docker Docs. "
    "https://docs.docker.com/desktop/features/wsl/",
    "Docker Inc. (2026c). Install Docker Engine on Ubuntu. Docker Docs. "
    "https://docs.docker.com/engine/install/ubuntu/",
    "Docker Inc. (2026d). Post-installation steps for Linux. Docker Docs. "
    "https://docs.docker.com/engine/install/linux-postinstall/",
    "Docker Inc. (2026e). Dockerfile reference. Docker Docs. "
    "https://docs.docker.com/reference/dockerfile/",
    "Docker Inc. (2026f). Build context. Docker Docs. "
    "https://docs.docker.com/build/concepts/context/",
    "Docker Inc. (2026g). Bind mounts. Docker Docs. "
    "https://docs.docker.com/engine/storage/bind-mounts/",
    "Docker Inc. (2026h). Environment variables in Compose. Docker Docs. "
    "https://docs.docker.com/compose/how-tos/environment-variables/",
    "Docker Inc. (2026i). Publishing and exposing ports. Docker Docs. "
    "https://docs.docker.com/get-started/docker-concepts/running-containers/publishing-ports/",
    "Docker Inc. (2026j). Run the Docker daemon as a non-root user (rootless mode). "
    "Docker Docs. https://docs.docker.com/engine/security/rootless/",
    "Servicio Nacional de Aprendizaje. (2026). Instalación y uso de Docker [Guía de "
    "aprendizaje]. Formación en Ambientes Virtuales de Aprendizaje.",
]
for r in sorted(refs):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    par.paragraph_format.first_line_indent = Cm(-1.27)
    par.paragraph_format.left_indent = Cm(1.27)
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run(r)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.save(OUT)
print("OK ->", OUT)
print("Figuras:", FIG_N[0], "| Tablas:", TABLE_N[0])
